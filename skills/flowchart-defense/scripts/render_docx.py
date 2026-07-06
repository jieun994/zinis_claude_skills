#!/usr/bin/env python3
# 검토리포트.md → 워드(.docx) 변환 (옵션, 고객 전달본용).
# - mermaid 코드블록은 mermaid-cli(mmdc)로 PNG 렌더해 워드에 이미지로 삽입.
# - mmdc가 없으면 mermaid를 코드블록 그대로 넣고 "흐름도는 MD에서 확인" 안내(작업 손실 0).
# - 마크다운 표는 워드 표로, 제목/문단은 그대로 변환(간단 파서 — 복잡 문법은 미지원).
# 사용: python render_docx.py <input.md> <output.docx>
# 사전: pip install python-docx  (렌더하려면 npm i -g @mermaid-js/mermaid-cli)
import sys
import os
import re
import subprocess
import shutil
import tempfile


def force_korean_font(doc, name="Malgun Gothic"):
    """모든 run에 동아시아(한글) 폰트를 강제 지정한다.
    python-docx는 기본적으로 w:eastAsia 폰트를 안 넣어서 한글이 □□□로 깨진다."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def fix_run(run):
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:eastAsia"), name)
        rfonts.set(qn("w:ascii"), name)
        rfonts.set(qn("w:hAnsi"), name)

    # 문서 기본 폰트(Normal)도 지정
    try:
        normal = doc.styles["Normal"]
        normal.font.name = name
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), name)
    except Exception:
        pass

    for p in doc.paragraphs:
        for r in p.runs:
            fix_run(r)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        fix_run(r)


def find_mmdc():
    for cmd in ("mmdc", "mmdc.cmd"):
        if shutil.which(cmd):
            return shutil.which(cmd)
    return None


def render_mermaid(code, out_png, mmdc):
    """mermaid 코드 → PNG. 성공 시 True."""
    with tempfile.NamedTemporaryFile("w", suffix=".mmd", delete=False, encoding="utf-8") as f:
        f.write(code)
        src = f.name
    try:
        subprocess.run(
            [mmdc, "-i", src, "-o", out_png, "-b", "white", "-s", "2"],
            check=True, capture_output=True, timeout=120,
        )
        return os.path.isfile(out_png)
    except Exception as e:  # noqa: BLE001
        print("mermaid render failed: %s" % e, file=sys.stderr)
        return False
    finally:
        try:
            os.unlink(src)
        except OSError:
            pass


def split_table(lines, i):
    """i번째 줄부터 마크다운 표를 읽어 (rows, next_i) 반환. 표가 아니면 (None, i)."""
    if "|" not in lines[i]:
        return None, i
    # 구분선(---|---) 확인
    if i + 1 >= len(lines) or not re.match(r"^\s*\|?[\s:|-]+\|[\s:|-]+", lines[i + 1]):
        return None, i
    rows = []
    j = i
    while j < len(lines) and "|" in lines[j]:
        if re.match(r"^\s*\|?[\s:|-]+\|[\s:|-]+\s*$", lines[j]):
            j += 1
            continue
        cells = [c.strip() for c in lines[j].strip().strip("|").split("|")]
        rows.append(cells)
        j += 1
    return rows, j


def main():
    if len(sys.argv) < 3:
        print("usage: render_docx.py <input.md> <output.docx>", file=sys.stderr)
        sys.exit(2)
    md_path, docx_path = sys.argv[1], sys.argv[2]
    if not os.path.isfile(md_path):
        print("input not found: %s" % md_path, file=sys.stderr)
        sys.exit(1)

    try:
        from docx import Document  # type: ignore
        from docx.shared import Inches
    except ImportError:
        print("python-docx 미설치 — 'pip install python-docx' 후 다시 실행하세요.", file=sys.stderr)
        sys.exit(1)

    mmdc = find_mmdc()
    if not mmdc:
        print("[안내] mermaid-cli(mmdc) 없음 — 흐름도는 코드블록으로 들어갑니다. "
              "이미지로 넣으려면 'npm i -g @mermaid-js/mermaid-cli' 설치 후 다시 실행하세요.",
              file=sys.stderr)

    text = open(md_path, encoding="utf-8").read()
    lines = text.splitlines()
    doc = Document()
    img_dir = tempfile.mkdtemp()
    img_n = 0

    i = 0
    while i < len(lines):
        line = lines[i]

        # mermaid 코드블록
        m = re.match(r"^```\s*mermaid\s*$", line.strip())
        if m:
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1  # closing ```
            code = "\n".join(block)
            rendered = False
            if mmdc:
                img_n += 1
                png = os.path.join(img_dir, "flow_%d.png" % img_n)
                if render_mermaid(code, png, mmdc):
                    doc.add_picture(png, width=Inches(6.0))
                    rendered = True
            if not rendered:
                doc.add_paragraph("[흐름도 — 아래 mermaid 코드. 렌더는 MD/뷰어에서 확인]")
                p = doc.add_paragraph(code)
                p.style = doc.styles["Normal"]
            continue

        # 기타 코드블록은 그대로
        if line.strip().startswith("```"):
            i += 1
            buf = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            doc.add_paragraph("\n".join(buf))
            continue

        # 제목
        h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if h:
            doc.add_heading(h.group(2).strip(), level=min(len(h.group(1)), 4))
            i += 1
            continue

        # 표
        rows, nj = split_table(lines, i)
        if rows:
            table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
            table.style = "Light Grid Accent 1"
            for ri, r in enumerate(rows):
                for ci, c in enumerate(r):
                    table.rows[ri].cells[ci].text = c
            i = nj
            doc.add_paragraph("")
            continue

        # 인용/일반 문단
        stripped = re.sub(r"^>\s?", "", line)
        if stripped.strip():
            doc.add_paragraph(stripped)
        else:
            doc.add_paragraph("")
        i += 1

    force_korean_font(doc)
    doc.save(docx_path)
    print("saved: %s%s" % (docx_path, "" if mmdc else "  (흐름도=코드블록, mmdc 미설치)"))


if __name__ == "__main__":
    main()
